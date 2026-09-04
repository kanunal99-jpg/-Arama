import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_CV_BYTES = 10 * 1024 * 1024

SKILL_ALIASES = {
    "python": "Python",
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "react": "React",
    "next.js": "Next.js",
    "nextjs": "Next.js",
    "node.js": "Node.js",
    "nodejs": "Node.js",
    "java": "Java",
    "c#": "C#",
    "sql": "SQL",
    "postgresql": "PostgreSQL",
    "mysql": "MySQL",
    "excel": "Microsoft Excel",
    "power bi": "Power BI",
    "sap": "SAP",
    "salesforce": "Salesforce",
    "docker": "Docker",
    "kubernetes": "Kubernetes",
    "aws": "AWS",
    "azure": "Azure",
    "git": "Git",
    "github": "GitHub",
    "scrum": "Scrum",
    "agile": "Agile",
    "project management": "Project Management",
    "sales": "Sales",
    "marketing": "Marketing",
    "leadership": "Leadership",
    "communication": "Communication",
}


def _extract_text(filename: str, data: bytes) -> str:
    extension = Path(filename or "").suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Desteklenmeyen CV formatı. PDF, DOCX veya TXT kullanın.")

    if extension == ".pdf":
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if extension == ".docx":
        document = Document(BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)

    return data.decode("utf-8", errors="replace")


def _first_match(pattern: str, text: str) -> str | None:
    match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
    return match.group(1).strip() if match else None


def _is_section_header(line: str, aliases: list[str]) -> bool:
    normalized = line.strip().lower().rstrip(":").strip()
    return any(normalized == alias or normalized.startswith(alias + " ") for alias in aliases)


def parse_cv(filename: str, data: bytes) -> dict:
    if not data:
        raise ValueError("CV dosyası boş.")
    if len(data) > MAX_CV_BYTES:
        raise ValueError("CV dosyası 10 MB sınırını aşıyor.")

    text = re.sub(r"[ \t]+", " ", _extract_text(filename, data)).strip()
    if not text:
        raise ValueError("CV metni çıkarılamadı. Metin tabanlı bir PDF/DOCX deneyin.")

    email = _first_match(r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", text)
    phone = _first_match(r"((?:\+?90[ .-]?)?(?:\(?\d{3}\)?[ .-]?)\d{3}[ .-]?\d{2}[ .-]?\d{2})", text)

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    name = None
    if lines:
        for line in lines[:8]:
            if len(line.split()) in (2, 3, 4) and not re.search(r"@|\d", line):
                name = line
                break

    lowered = text.lower()
    skills = []
    for alias, canonical in SKILL_ALIASES.items():
        if re.search(rf"(?<!\w){re.escape(alias)}(?!\w)", lowered):
            skills.append(canonical)

    section_names = {
        "experience": [
            "experience",
            "work experience",
            "iş deneyimi",
            "deneyim",
            "mesleki deneyim",
        ],
        "education": [
            "education",
            "eğitim",
            "eğitim bilgileri",
        ],
        "skills": [
            "skills",
            "yetenekler",
            "beceriler",
            "yetkinlikler",
        ],
    }
    boundary_sections = {
        **section_names,
        "references": ["references", "referanslar"],
    }

    sections = {}
    lower_lines = [line.lower().strip() for line in lines]
    for key, aliases in section_names.items():
        indexes = [
            i for i, line in enumerate(lower_lines)
            if _is_section_header(line, aliases)
        ]
        if not indexes:
            continue

        start = indexes[0] + 1
        end = len(lines)
        for i in range(start, len(lines)):
            if any(_is_section_header(lower_lines[i], values) for values in boundary_sections.values()):
                end = i
                break

        content = lines[start:end][:50]
        sections[key] = content

        # Turkish CVs commonly use a dedicated YETKİNLİKLER section with
        # human-readable skills that are not present in the English alias map.
        if key == "skills":
            section_skills = []
            for item in content:
                cleaned = re.sub(r"^[•●▪◦*-]+\s*", "", item).strip()
                if cleaned and cleaned not in section_skills:
                    section_skills.append(cleaned)
            if section_skills:
                skills = section_skills

    return {
        "filename": filename,
        "text_length": len(text),
        "profile": {
            "name": name,
            "email": email,
            "phone": phone,
            "skills": skills,
            "sections": sections,
        },
    }
